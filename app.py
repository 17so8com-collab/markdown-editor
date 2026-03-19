#!/usr/bin/env python3
import os
import sys
import webview
import markdown
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


class Api:
    def __init__(self):
        self.current_file = None
        self.is_modified = False

    def read_file_dialog(self):
        import tkinter as tk
        from tkinter import filedialog

        root = tk.Tk()
        root.withdraw()
        root.attributes("-topmost", True)
        path = filedialog.askopenfilename(
            title="打开文件",
            filetypes=[
                ("所有支持", "*.md *.markdown *.docx *.html *.htm"),
                ("Markdown", "*.md *.markdown"),
                ("Word文档", "*.docx"),
                ("HTML文件", "*.html *.htm"),
                ("所有文件", "*.*"),
            ],
        )
        root.destroy()
        if not path:
            return json.dumps(None)
        try:
            if path.endswith(".docx"):
                result = self.docx_to_md(path)
                if result["success"]:
                    self.current_file = path
                    return json.dumps(
                        {"success": True, "content": result["content"], "path": path}
                    )
                else:
                    return json.dumps({"success": False, "error": result["error"]})
            elif path.endswith((".html", ".htm")):
                with open(path, "r", encoding="utf-8") as f:
                    html_content = f.read()
                result = self.html_to_md(html_content)
                if result["success"]:
                    self.current_file = path
                    return json.dumps(
                        {"success": True, "content": result["content"], "path": path}
                    )
                else:
                    return json.dumps({"success": False, "error": result["error"]})
            else:
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read()
                self.current_file = path
                return json.dumps({"success": True, "content": content, "path": path})
        except Exception as e:
            return json.dumps({"success": False, "error": str(e)})

    def save_file_dialog(self):
        import tkinter as tk
        from tkinter import filedialog

        root = tk.Tk()
        root.withdraw()
        root.attributes("-topmost", True)
        path = filedialog.asksaveasfilename(
            title="保存文件",
            defaultextension=".md",
            filetypes=[
                ("Markdown", "*.md"),
                ("Markdown", "*.markdown"),
                ("Word文档", "*.docx"),
                ("HTML文件", "*.html"),
                ("所有文件", "*.*"),
            ],
        )
        root.destroy()
        if not path:
            return json.dumps(None)
        return json.dumps(path)

    def save_file(self, path, content):
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
            self.current_file = path
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def get_current_file(self):
        return self.current_file

    def md_to_html(self, md_content):
        try:
            html = markdown.markdown(
                md_content, extensions=["tables", "fenced_code", "codehilite", "toc"]
            )
            return {"success": True, "html": html}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def md_to_docx(self, md_content, output_path):
        try:
            doc = Document()
            doc.add_heading("Markdown Document", 0)

            lines = md_content.split("\n")
            in_code_block = False
            code_buffer = []

            for line in lines:
                if line.startswith("```"):
                    if in_code_block:
                        p = doc.add_paragraph()
                        run = p.add_run("\n".join(code_buffer))
                        run.font.name = "Consolas"
                        run.font.size = Pt(10)
                        code_buffer = []
                    else:
                        in_code_block = True
                    continue

                if in_code_block:
                    code_buffer.append(line)
                    continue

                if line.startswith("# "):
                    doc.add_heading(line[2:], 1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], 2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], 3)
                elif line.startswith("- ") or line.startswith("* "):
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(line[2:])
                elif line.startswith("> "):
                    p = doc.add_paragraph(style="Quote")
                    p.add_run(line[2:])
                elif line.strip() == "":
                    doc.add_paragraph()
                else:
                    text = line
                    p = doc.add_paragraph()
                    while "`" in text:
                        idx = text.find("`")
                        if idx > 0:
                            p.add_run(text[:idx])
                        text = text[idx + 1 :]
                        end = text.find("`")
                        if end == -1:
                            p.add_run("`" + text)
                            break
                        run = p.add_run(text[:end])
                        run.font.name = "Consolas"
                        run.font.size = Pt(10)
                        text = text[end + 1 :]
                    if text:
                        p.add_run(text)

            doc.save(output_path)
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def docx_to_md(self, file_path):
        try:
            doc = Document(file_path)
            md_lines = []
            style_name = (
                doc.paragraphs[0].style.name
                if doc.paragraphs and doc.paragraphs[0].style
                else ""
            )
            md_lines.append(
                "# "
                + (
                    doc.paragraphs[0].text
                    if doc.paragraphs and doc.paragraphs[0].text
                    else "Document"
                )
            )
            md_lines.append("")

            for para in doc.paragraphs[1:]:
                text = para.text.strip()
                if not text:
                    md_lines.append("")
                    continue
                s_name = para.style.name if para.style else ""
                if "Heading" in s_name:
                    level = s_name.split(" ")[-1]
                    try:
                        level = int(level)
                        md_lines.append("#" * level + " " + text)
                    except (ValueError, AttributeError):
                        md_lines.append("## " + text)
                elif "Quote" in s_name:
                    md_lines.append("> " + text)
                else:
                    md_lines.append(text)
                md_lines.append("")

            return {"success": True, "content": "\n".join(md_lines)}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def html_to_md(self, html_content):
        try:
            import re

            md = html_content

            for i in range(6, 0, -1):
                md = re.sub(
                    r"<h" + str(i) + r"[^>]*>(.*?)</h" + str(i) + r">",
                    lambda m: "#" * i
                    + " "
                    + re.sub(r"<[^>]+>", "", m.group(1))
                    + "\n\n",
                    md,
                    flags=re.DOTALL,
                )

            md = re.sub(
                r"<p[^>]*>(.*?)</p>",
                lambda m: re.sub(r"<[^>]+>", "", m.group(1)) + "\n\n",
                md,
                flags=re.DOTALL,
            )

            md = re.sub(
                r"<pre[^>]*><code[^>]*>(.*?)</code></pre>",
                lambda m: "```\n" + re.sub(r"<[^>]+>", "", m.group(1)) + "\n```\n\n",
                md,
                flags=re.DOTALL,
            )

            md = re.sub(r"<code[^>]*>(.*?)</code>", r"`\1`", md, flags=re.DOTALL)
            md = re.sub(r"<(strong|b)[^>]*>(.*?)</\1>", r"**\2**", md, flags=re.DOTALL)
            md = re.sub(r"<(em|i)[^>]*>(.*?)</\1>", r"*\2*", md, flags=re.DOTALL)
            md = re.sub(
                r'<a[^>]*href=["\']([^"\']+)["\'][^>]*>(.*?)</a>',
                r"[\2](\1)",
                md,
                flags=re.DOTALL,
            )
            md = re.sub(
                r"<li[^>]*>(.*?)</li>",
                lambda m: "- " + re.sub(r"<[^>]+>", "", m.group(1)) + "\n",
                md,
                flags=re.DOTALL,
            )
            md = re.sub(r"<ul[^>]*>|</ul>", "\n", md, flags=re.DOTALL)
            md = re.sub(r"<[^>]+>", "", md)
            md = re.sub(r"\n{3,}", "\n\n", md)

            return {"success": True, "content": md.strip()}
        except Exception as e:
            return {"success": False, "error": str(e)}


def create_window():
    api = Api()
    window = webview.create_window(
        "Markdown Editor",
        html=_get_html(),
        width=1400,
        height=900,
        min_size=(800, 600),
        resizable=True,
        js_api=api,
    )
    webview.start()


def _get_html():
    return """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1"/>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/vditor@3.11.0/dist/index.css"/>
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,sans-serif;overflow:hidden;height:100vh;display:flex;flex-direction:column}
.toolbar{display:flex;align-items:center;gap:8px;padding:8px 16px;background:#f5f5f5;border-bottom:1px solid #ddd;flex-shrink:0;flex-wrap:wrap}
.toolbar.dark-mode{background:#1e1e1e;border-color:#333}
.toolbar-group{display:flex;align-items:center;gap:4px}
.toolbar-sep{width:1px;height:24px;background:#ccc;margin:0 4px}
.toolbar.dark-mode .toolbar-sep{background:#444}
.toolbar button,.toolbar select{padding:6px 12px;border:1px solid #ddd;border-radius:4px;background:#fff;cursor:pointer;font-size:13px;white-space:nowrap}
.toolbar button:hover,.toolbar select:hover{background:#e8e8e8}
.toolbar button.active{background:#4a90d9;color:#fff;border-color:#4a90d9}
.toolbar.dark-mode button,.toolbar.dark-mode select{background:#2d2d2d;border-color:#444;color:#ccc}
.toolbar.dark-mode button:hover,.toolbar.dark-mode select:hover{background:#3d3d3d}
.toolbar .title{min-width:200px;text-align:center;font-weight:bold;color:#666;font-size:13px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
.toolbar.dark-mode .title{color:#aaa}
.toolbar .logo{font-weight:bold;color:#4a90d9;font-size:14px}
.editor-container{flex:1;overflow:hidden}
#vditor{height:100%}
.statusbar{display:flex;align-items:center;gap:16px;padding:4px 16px;background:#f5f5f5;border-top:1px solid #ddd;font-size:12px;color:#666;flex-shrink:0}
.statusbar.dark-mode{background:#1e1e1e;border-color:#333;color:#999}
.statusbar span{display:flex;align-items:center;gap:4px}
.statusbar .dot{width:8px;height:8px;border-radius:50%;background:#4a90d9}
.statusbar .dot.modified{background:#f5a623}
.modal-overlay{position:fixed;inset:0;background:rgba(0,0,0,0.5);display:none;align-items:center;justify-content:center;z-index:1000}
.modal-overlay.show{display:flex}
.modal{background:#fff;border-radius:8px;padding:24px;min-width:400px;max-width:600px;box-shadow:0 4px 24px rgba(0,0,0,0.2)}
.modal h3{font-size:16px;margin-bottom:16px;color:#333}
.modal p{margin-bottom:16px;color:#666;font-size:13px;line-height:1.5}
.modal-actions{display:flex;gap:8px;justify-content:flex-end}
.modal button{padding:8px 20px;border-radius:4px;border:1px solid #ddd;background:#fff;cursor:pointer;font-size:13px}
.modal button.primary{background:#4a90d9;color:#fff;border-color:#4a90d9}
.modal button:hover{opacity:0.9}
::-webkit-scrollbar{width:8px;height:8px}
::-webkit-scrollbar-track{background:#f0f0f0}
::-webkit-scrollbar-thumb{background:#ccc;border-radius:4px}
::-webkit-scrollbar-thumb:hover{background:#aaa}
.dark-mode ::-webkit-scrollbar-track{background:#1e1e1e}
.dark-mode ::-webkit-scrollbar-thumb{background:#444}

/* Vditor editor dark mode - applied when .vditor--dark class exists */
.vditor--dark{background:#1e1e1e}
.vditor--dark .vditor-toolbar{background:#252526;border-color:#3c3c3c}
.vditor--dark .vditor-toolbar__icon{color:#ccc}
.vditor--dark .vditor-toolbar__icon:hover{background:#3c3c3c}
.vditor--dark .vditor-toolbar__icon.vditor-menu--current{background:#3c3c3c}
.vditor--dark .vditor-wysiwyg{background:#252526}
.vditor--dark .vditor-wysiwyg .vditor-reset{color:#d4d4d4}
.vditor--dark .vditor-wysiwyg .vditor-reset h1,
.vditor--dark .vditor-wysiwyg .vditor-reset h2,
.vditor--dark .vditor-wysiwyg .vditor-reset h3,
.vditor--dark .vditor-wysiwyg .vditor-reset h4,
.vditor--dark .vditor-wysiwyg .vditor-reset h5,
.vditor--dark .vditor-wysiwyg .vditor-reset h6{color:#d4d4d4}
.vditor--dark .vditor-wysiwyg .vditor-reset p{color:#d4d4d4}
.vditor--dark .vditor-wysiwyg .vditor-reset pre{background:#2d2d2d}
.vditor--dark .vditor-wysiwyg .vditor-reset pre code{color:#ce9178}
.vditor--dark .vditor-wysiwyg .vditor-reset blockquote{color:#aaa;border-color:#555}
.vditor--dark .vditor-wysiwyg .vditor-reset table{color:#d4d4d4}
.vditor--dark .vditor-wysiwyg .vditor-reset table th,
.vditor--dark .vditor-wysiwyg .vditor-reset table td{border-color:#3c3c3c}
.vditor--dark .vditor-wysiwyg .vditor-reset table th{background:#252526}
.vditor--dark .vditor-wysiwyg .vditor-reset a{color:#6db3f2}
.vditor--dark .vditor-wysiwyg .vditor-reset hr{border-color:#3c3c3c}
.vditor--dark .vditor-ir{background:#252526}
.vditor--dark .vditor-ir .vditor-reset{color:#d4d4d4}
.vditor--dark .vditor-sv{background:#252526}
.vditor--dark .vditor-sv .vditor-reset{color:#d4d4d4}
.vditor--dark .vditor-preview{background:#252526}
.vditor--dark .vditor-preview .vditor-reset{color:#d4d4d4;background:#252526}
.vditor--dark .vditor-hint{background:#252526;border-color:#3c3c3c}
.vditor--dark .vditor-hint__item{color:#ccc}
.vditor--dark .vditor-hint__item:hover{background:#3c3c3c}
.vditor--dark .vditor-panel{background:#252526;border-color:#3c3c3c}
.vditor--dark .vditor-counter{color:#858585}
.vditor--dark .vditor-resume{color:#858585}
.vditor--dark .vditor-toc{background:#252526;border-color:#3c3c3c}
</style>
</head>
<body>
<div class="toolbar" id="toolbar">
  <span class="logo">MD</span>
  <div class="toolbar-group">
    <button onclick="newFile()" title="新建">新建</button>
    <button onclick="openFile()" title="打开">打开</button>
    <button onclick="saveFile()" title="保存">保存</button>
    <button onclick="saveFileAs()" title="另存为">另存为</button>
  </div>
  <div class="toolbar-sep"></div>
  <div class="toolbar-group">
    <button onclick="togglePreview()" id="previewBtn" title="切换预览">预览</button>
    <button onclick="setFullscreen()" id="fullscreenBtn" title="全屏">全屏</button>
  </div>
  <div class="toolbar-sep"></div>
  <div class="toolbar-group">
    <button onclick="exportHTML()" title="导出HTML">HTML</button>
    <button onclick="exportDOCX()" title="导出Word">Word</button>
    <button onclick="exportPDF()" title="导出PDF">PDF</button>
    <button onclick="exportMarkdown()" title="导出Markdown">MD</button>
  </div>
  <div class="toolbar-sep"></div>
  <button onclick="toggleTheme()" id="themeBtn" title="切换主题">深色</button>
  <span class="title" id="titleBar">未命名.md</span>
</div>
<div class="editor-container"><div id="vditor"></div></div>
<div class="statusbar" id="statusbar">
  <span><span class="dot" id="statusDot"></span><span id="statusText">就绪</span></span>
  <span id="wordCount">字数: 0</span>
  <span id="charCount">字符: 0</span>
  <span id="lineCount">行数: 0</span>
</div>

<div class="modal-overlay" id="unsavedModal">
  <div class="modal">
    <h3>未保存的更改</h3>
    <p id="modalMsg">当前文件有未保存的更改，是否保存？</p>
    <div class="modal-actions">
      <button onclick="closeModal()">取消</button>
      <button onclick="closeModal();discardChanges()">不保存</button>
      <button class="primary" onclick="closeModal();confirmSave()">保存</button>
    </div>
  </div>
</div>

<script src="https://cdn.jsdelivr.net/npm/vditor@3.11.0/dist/index.min.js"></script>
<script>
let vditor;
let currentPath = null;
let isModified = false;
let pendingAction = null;
let isDarkMode = false;

function newFile(){checkAndDo(function(){vditor.setValue("# \\u65b0\\u6587\\u6863\\n\\n");currentPath=null;setTitle("\\u672a\\u547d\\u540d.md");setModified(false);updateStatus("\\u65b0\\u5efa\\u6587\\u6863");});}
function openFile(){window.pywebview.api.read_file&&pythonOpenFile();}
async function pythonOpenFile(){var result=await window.pywebview.api.read_file_dialog();if(result.returnValue){var resp=JSON.parse(result.returnValue);if(resp.success){vditor.setValue(resp.content);currentPath=resp.path||null;var name=currentPath?currentPath.split(/[\\\\\\/]/).pop():"\\u6253\\u5f00\\u7684\\u6587\\u4ef6";setTitle(name);setModified(false);updateStatus("\\u5df2\\u6253\\u5f00: "+name);}else{alert("\\u6253\\u5f00\\u5931\\u6548: "+resp.error);}}}
function saveFile(){if(currentPath){doSave(currentPath);}else{saveFileAs();}}
async function saveFileAs(){var result=await window.pywebview.api.save_file_dialog();if(result.returnValue){var path=JSON.parse(result.returnValue);if(path){doSave(path);}}}
async function doSave(path){var content=vditor.getValue();var result=await window.pywebview.api.save_file(path,content);if(result.returnValue){var resp=JSON.parse(result.returnValue);if(resp.success){currentPath=path;var name=path.split(/[\\\\\\/]/).pop();setTitle(name);setModified(false);updateStatus("\\u5df2\\u4fdd\\u5b58: "+name);}else{alert("\\u4fdd\\u5b58\\u5931\\u6548: "+resp.error);}}}
function checkAndDo(action){if(isModified){pendingAction=action;document.getElementById("unsavedModal").classList.add("show");}else{action&&action();}}
function closeModal(){document.getElementById("unsavedModal").classList.remove("show");}
function discardChanges(){pendingAction&&pendingAction();pendingAction=null;}
async function confirmSave(){await saveFile();pendingAction&&pendingAction();pendingAction=null;}
function setTitle(name){document.getElementById("titleBar").textContent=name;}
function setModified(mod){isModified=mod;var dot=document.getElementById("statusDot");if(mod){dot.classList.add("modified");}else{dot.classList.remove("modified");}}
function updateStatus(text){document.getElementById("statusText").textContent=text;}
function togglePreview(){vditor.togglePreview&&vditor.togglePreview();}
function setFullscreen(){if(document.fullscreenElement){document.exitFullscreen();}else{document.body.requestFullscreen();}}
async function exportMarkdown(){var result=await window.pywebview.api.save_file_dialog();if(result.returnValue){var path=JSON.parse(result.returnValue);if(path){var content=vditor.getValue();await window.pywebview.api.save_file(path,content);updateStatus("\\u5df2\\u5bfc\\u51fa: "+path.split(/[\\\\\\/]/).pop());}}}
async function exportHTML(){var result=await window.pywebview.api.save_file_dialog();if(result.returnValue){var path=JSON.parse(result.returnValue);if(path){var html=vditor.getHTML();var fullHtml="<!DOCTYPE html><html><head><meta charset=\\"utf-8\\"><title>Markdown Export</title><style>body{max-width:800px;margin:40px auto;padding:20px;font-family:sans-serif;line-height:1.8;}code{background:#f5f5f5;padding:2px 6px;border-radius:3px;}.vditor-reset pre{background:#f5f5f5;padding:16px;overflow:auto;border-radius:6px;}</style></head><body>"+html+"</body></html>";await window.pywebview.api.save_file(path,fullHtml);updateStatus("\\u5df2\\u5bfc\\u51faHTML");}}}
async function exportDOCX(){var result=await window.pywebview.api.save_file_dialog();if(result.returnValue){var path=JSON.parse(result.returnValue);if(path){var content=vditor.getValue();await window.pywebview.api.md_to_docx(content,path);updateStatus("\\u5df2\\u5bfc\\u51faWord");}}}
function exportPDF(){window.print();}
function toggleTheme(){isDarkMode=!isDarkMode;var btn=document.getElementById("themeBtn");var toolbar=document.getElementById("toolbar");var statusbar=document.getElementById("statusbar");if(isDarkMode){btn.textContent="\\u6d45\\u8272";toolbar.classList.add("dark-mode");statusbar.classList.add("dark-mode");document.body.style.background="#1e1e1e";Vditor.setTheme("dark","dark");}else{btn.textContent="\\u6df1\\u8272";toolbar.classList.remove("dark-mode");statusbar.classList.remove("dark-mode");document.body.style.background="#fff";Vditor.setTheme("classic","light");}}
function updateCounts(){var content=vditor.getValue();var chars=content.length;var words=content.replace(/[\\n\\s]+/g," ").split(" ").filter(function(w){return w;}).length;var lines=content.split("\\n").length;document.getElementById("wordCount").textContent="\\u5b57\\u6570: "+words;document.getElementById("charCount").textContent="\\u5b57\\u7b26: "+chars;document.getElementById("lineCount").textContent="\\u884c\\u6570: "+lines;}

window.addEventListener("load",function(){
  vditor=new Vditor("vditor",{
    height:"100%",
    mode:"wysiwyg",
    placeholder:"\\u5f00\\u59cb\\u5199\\u4f5c...",
    typewriterMode:true,
    tab:"    ",
    toolbarConfig:{pin:true},
    toolbar:[
      "headings",
      "bold",
      "italic",
      "strike",
      "|",
      "line",
      "quote",
      "list",
      "ordered-list",
      "check",
      "outdent",
      "indent",
      "|",
      "code",
      "inline-code",
      "link",
      "table",
      "|",
      "undo",
      "redo",
      "|",
      "preview",
      "fullscreen",
      "edit-mode"
    ],
    input:function(){setModified(true);updateCounts();},
    after:function(){updateCounts();}
  });
});
</script>
</body>
</html>"""


if __name__ == "__main__":
    create_window()
